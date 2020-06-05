#!/usr/bin/env python
# coding: utf-8

# In[1]:

bookname = input('Put book/movie name: ')
bookname = '_'.join(bookname.split())

# In[2]:

import pandas as pd
import sys


import os
import os.path
from os import path

from datetime import datetime

newflag = False
if not os.path.exists(bookname):
    newflag = True
    print('[{}] Create folders and files that you need'.format(datetime.now().time()))
    os.makedirs(bookname)
    pathprefix = './'+bookname+'/'
    files = [pathprefix+bookname+'_new.xslx', pathprefix+bookname+'_DB.csv']
    nondf = pd.DataFrame()
    nondf.to_csv('./'+bookname+'/'+bookname+'_DB.csv', index = False)
    nondf.to_excel('./' + bookname + '/' + bookname + '_new.xlsx', index=False)
    f = open('./' + bookname + '/' + bookname + '_DB.csv', 'w')
    f.close()
    print('[{}] New book/movie folder is generated. Select the languages that you\'re gonna use...'.format(datetime.now().time()))
    srclang = input('Enter Source Language (Learning Language): ')
    tgtlang = input('Enter Target Language (Native Language): ')
    srclang = srclang[0].upper()+srclang[1:]
    tgtlang = tgtlang[0].upper() + tgtlang[1:]

# In[3]:

try:
    DB=pd.read_csv('./'+bookname+'/'+bookname+'_DB.csv')
    print('[{}] Loading DB file complete'.format(datetime.now().time(), bookname))
    srclang = DB.columns[0]
    tgtlang = DB.columns[1]

except:
    if not newflag:
        print('[{}] Oops! {} DB seems deleted'.format(datetime.now().time(), bookname))
        print('Select the languages that you\'re gonna use to build a DB again..')
        srclang = input('Enter Source Language (Learning Language): ')
        tgtlang = input('Enter Target Language (Native Language): ')
        srclang = srclang[0].upper() + srclang[1:]
        tgtlang = tgtlang[0].upper() + tgtlang[1:]

    if not os.path.exists('./' + bookname + '/' + bookname + '_DB.csv'):
        print('{}_DB.csv file doesn\'t exist'.format(bookname))
        print('[{}] Creating {}_DB.csv file...'.format(datetime.now().time(), bookname))
    f = open('./'+bookname+'/'+bookname+'_DB.csv' , 'w')
    DB = pd.DataFrame(columns = [srclang,tgtlang,'cnt'])
    DB.to_csv('./' + bookname + '/' + bookname + '_DB.csv', index=False, encoding='utf-8-sig')

    f.close()


try:
    newdata = pd.read_excel('./'+bookname+'/'+bookname+'_new.xlsx',header = None, names = ['src','tgt'], sep=",\t")
    print('[{}] Reading {}_new.xlsx file complete'.format(datetime.now().time(), bookname))
except:
    if not os.path.exists('./'+bookname+'/'+bookname+'_new.xlsx'):
        print('{}_new.xlsx file doesn\'t exist or empty'.format(bookname))
        print('[{}] Creating {}_new.xlsx file...'.format(datetime.now().time(), bookname))
    #f=open('./' + bookname + '/' + bookname + '_new.xslx','w')
    nondf = pd.DataFrame()
    nondf.to_excel('./' + bookname + '/' + bookname + '_new.xlsx', index=False)
    newdata = pd.DataFrame(columns = ['src','tgt'])

if newdata.empty:
    print('[{}] {}_new.xlsx file is empty!'.format(datetime.now().time(), bookname))
    print('Please put words in {}_new.xlsx file!'.format(bookname))
    sys.exit()


# In[4]:


print('[{}] Building Dictionary...'.format(datetime.now().time()))

from collections import Counter
srcwdict = dict()
srcwcounter = Counter()
for i, row in DB.iterrows():
    srcwdict[row[srclang]] = row[tgtlang]
    srcwcounter[row[srclang]] = row.cnt


# In[5]:


seen = set()
newwords = set(newdata.src)
for i, row in newdata.iterrows():
    if row.src in srcwdict:
        srcwcounter[row.src] += 1
        seen.add(row.src)
    else:
        srcwdict[row.src] = row.tgt
        srcwcounter[row.src] = 1


# In[6]:


srclst = []
tgtlst = []
cntlst = []
for srcw, tgtw in srcwdict.items():
    srclst.append(srcw)
    tgtlst.append(tgtw)
    cntlst.append(srcwcounter[srcw])


# In[7]:


data = pd.DataFrame({srclang: srclst, tgtlang: tgtlst,'cnt': cntlst})




# In[11]:


# In[13]:


from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import RGBColor
import tqdm
 
print('[{}] Writing on Vocabulary MS Words file...'.format(datetime.now().time()))
    
document = Document()

height = 0.3

print('[{}] newly added words part...'.format(datetime.now().time()))
 
document.add_heading(bookname+' Vocabulary', 0)
data = data.sort_values('cnt', ascending=False)

p = document.add_heading('Newly added words',level=1)

table_new = document.add_table(newdata.shape[0]+1, newdata.shape[1])
hdr_cells = table_new.rows[0].cells

languages = [srclang, tgtlang]
for j in range(newdata.shape[1]):
    run = hdr_cells[j].paragraphs[0].add_run(languages[j])
    run.bold = True
    run.underline = True

for i in tqdm.tqdm(range(newdata.shape[0])):
    for j in range(newdata.shape[-1]):
        cell = table_new.cell(i+1,j)
        cell.text = str(newdata.values[i,j])
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

for row in table_new.rows:
    row.height = Cm(height)

print('[{}] most frequently confusing words part...'.format(datetime.now().time()))
p = document.add_heading('Frequently confusing words',level=1)

cntMoreThanOne = sum(data.cnt > 1)
table_fq = document.add_table(cntMoreThanOne+1, data.shape[1])

hdr_cells = table_fq.rows[0].cells

# add the header rows.
for j in range(data.shape[-1]):
    run = hdr_cells[j].paragraphs[0].add_run(data.columns[j])
    run.bold = True
    run.underline = True
    #table_fq.cell(0,j).text = data.columns[j]

# add the rest of the data frame
seenbefore = False
#newreg = False
idx = 0
for i in tqdm.tqdm(range(data.shape[0])):
    if data.values[i,2] < 2: continue
    seenbefore = False
    #newreg = False
    if str(data.values[i,0]) in seen:
        seenbefore = True
    #if str(data.values[i,0]) in newwords:
    #    newreg = True
    for j in range(data.shape[-1]):
        cell = table_fq.cell(idx+1,j)
        cell.text = str(data.values[i,j])
        if seenbefore:
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        #elif newreg:
        #    run = cell.paragraphs[0].runs[0]
        #    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    idx += 1
        

for row in table_fq.rows:
    row.height = Cm(height)


data["lower"] = data[srclang].str.lower()
data.sort_values(['lower'], axis=0, ascending=True, inplace=True)
del data['lower']
save_data = data.copy()
del data['cnt']

# In[14]:


'''
Chosung extractor
Get code from https://frhyme.github.io/python/python_korean_englished/ and modify
'''

# 초성 리스트. 00 ~ 18
CHOSUNG_LIST = ['ㄱ', 'ㄲ', 'ㄴ', 'ㄷ', 'ㄸ', 'ㄹ', 'ㅁ', 'ㅂ', 'ㅃ', 'ㅅ', 'ㅆ', 'ㅇ', 'ㅈ', 'ㅉ', 'ㅊ', 'ㅋ', 'ㅌ', 'ㅍ', 'ㅎ']


def get_first_letter(word):
    r_lst = []
    w = word[0]
        ## 영어인 경우 구분해서 작성함.
    if '가' <= w <= '힣':
        ## 588개 마다 초성이 바뀜.
        ch1 = (ord(w) - ord('가')) // 588
        return CHOSUNG_LIST[ch1]
    else:
        return w

data['startletter'] = list(map(get_first_letter, data[srclang].str[0].str.lower()))
print('[{}] word dictionary part...'.format(datetime.now().time()))
p = document.add_heading('Word Dictionary',level=1)
data_group_letter = data.groupby('startletter')
existletters = sorted(data_group_letter.groups.keys())
for ch in existletters:
    #ch = chr(ord('a')+i)
    if ch in data_group_letter.groups.keys():
        p = document.add_heading(ch.upper(),level=2)
        groupdata = data_group_letter.get_group(ch)

        table_dict = document.add_table(groupdata.shape[0], groupdata.shape[1]-1)

        # add the rest of the data frame
        for i in range(groupdata.shape[0]):
            for j in range(groupdata.shape[-1]-1):
                table_dict.cell(i,j).text = str(groupdata.values[i,j])
        
        for row in table_dict.rows:
            row.height = Cm(height)



try:
    document.save('./'+bookname+'/'+bookname+'_Dictionary.docx') # Save document
    print('[{}] Document is created!!'.format(datetime.now().time()))
except:
    print('Word file is opened. Please close and re-run the program')
    sys.exit()

print('[{}] Saving new DB...'.format(datetime.now().time()))
save_data.to_csv('./'+bookname+'/'+bookname+'_DB.csv',index=False, encoding='utf-8-sig')


# In[9]:


nondf = pd.DataFrame()


# In[10]:


print('[{}] Clearing {}_new.xlsx file...'.format(datetime.now().time(), bookname))
nondf.to_excel('./' + bookname + '/' + bookname + '_new.xlsx', index=False)


# In[ ]:




